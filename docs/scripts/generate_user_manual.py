"""Generate Korean user manual (Word) for ER Modeler and DB 표준 점검 도구."""

from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[2]
IMG_DIR = ROOT / "docs" / "manual" / "images"
OUT_PATH = ROOT / "docs" / "manual" / "ER-Modeler-DB표준점검-사용자매뉴얼.docx"


def set_doc_defaults(doc: Document) -> None:
    style = doc.styles["Normal"]
    style.font.name = "맑은 고딕"
    style.font.size = Pt(11)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")


def add_toc(doc: Document) -> None:
    p = doc.add_paragraph()
    run = p.add_run()
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), 'TOC \\o "1-3" \\h \\z \\u')
    run._r.append(fld)
    doc.add_page_break()


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = "맑은 고딕"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")


def add_para(doc: Document, text: str, bold: bool = False) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "맑은 고딕"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(item, style="List Bullet")
        for run in p.runs:
            run.font.name = "맑은 고딕"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")


def add_numbered(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(item, style="List Number")
        for run in p.runs:
            run.font.name = "맑은 고딕"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")


def add_image(doc: Document, filename: str, caption: str, width_cm: float = 15.5) -> None:
    path = IMG_DIR / filename
    if not path.exists():
        add_para(doc, f"[이미지 없음: {filename}]", bold=True)
        return
    doc.add_picture(str(path), width=Cm(width_cm))
    cap = doc.add_paragraph(caption)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in cap.runs:
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(80, 80, 80)
        run.font.name = "맑은 고딕"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
    doc.add_paragraph()


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
    for r_idx, row in enumerate(rows, start=1):
        for c_idx, val in enumerate(row):
            table.rows[r_idx].cells[c_idx].text = val


def build() -> None:
    doc = Document()
    set_doc_defaults(doc)

    # Cover
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = title.add_run("MyPlatform\n사용자 매뉴얼")
    tr.bold = True
    tr.font.size = Pt(28)
    tr.font.name = "맑은 고딕"
    tr._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = sub.add_run("ER Modeler · DB 표준 점검 도구")
    sr.font.size = Pt(16)
    sr.font.name = "맑은 고딕"
    sr._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")

    ver = doc.add_paragraph()
    ver.alignment = WD_ALIGN_PARAGRAPH.CENTER
    vr = ver.add_run(f"문서 버전 1.0 · {date.today().isoformat()}")
    vr.font.size = Pt(11)
    vr.font.color.rgb = RGBColor(100, 100, 100)

    doc.add_page_break()

    add_heading(doc, "목차", 1)
    add_toc(doc)

    # 1. Introduction
    add_heading(doc, "1. 개요", 1)
    add_para(
        doc,
        "MyPlatform은 데이터베이스 설계·표준 점검 업무를 지원하는 웹 포털입니다. "
        "본 매뉴얼은 다음 두 애플리케이션의 사용 방법을 설명합니다.",
    )
    add_bullets(
        doc,
        [
            "ER Modeler — 테이블정의서 또는 SQL에서 ERD를 만들고 편집한 뒤 설계서·스크립트로 내보냅니다.",
            "DB 표준 점검 도구 — 행정안전부 공통표준(단어·용어·도메인·코드) 기준으로 설계서를 점검합니다.",
        ],
    )
    add_image(doc, "manual-workflow-overview.png", "그림 1-1. MyPlatform DB 설계·점검 업무 흐름")

    add_heading(doc, "1.1 사전 준비", 2)
    add_bullets(
        doc,
        [
            "포털(Next.js)과 API 서버(FastAPI, 기본 포트 8001)가 실행 중이어야 합니다.",
            "브라우저: Chrome, Edge 등 최신 브라우저 권장.",
            "ER Modeler: Excel 테이블정의서 또는 CREATE TABLE SQL.",
            "DB 표준 점검: 테이블정의서 Excel(단어·용어·도메인), 코드정의서 Excel(표준코드).",
        ],
    )

    add_heading(doc, "1.2 포털 접속", 2)
    add_numbered(
        doc,
        [
            "브라우저에서 MyPlatform 포털 주소에 접속합니다. (로컬 개발: http://localhost:3000)",
            "상단 메뉴 또는 앱 목록에서 원하는 앱을 선택합니다.",
            "ER Modeler: /apps/er-modeler",
            "DB 표준 점검 도구: /apps/chk-db-std",
        ],
    )

    # 2. ER Modeler
    doc.add_page_break()
    add_heading(doc, "2. ER Modeler", 1)
    add_para(
        doc,
        "ER Modeler는 테이블·컬럼·관계를 시각적으로 편집하는 ERD 도구입니다. "
        "Excel 정의서나 SQL을 가져와 다이어그램을 만들고, 검증 후 다시 Excel 또는 DDL 스크립트로 내보낼 수 있습니다.",
    )
    add_image(doc, "manual-er-modeler-ui.png", "그림 2-1. ER Modeler 화면 구성")

    add_heading(doc, "2.1 화면 구성", 2)
    add_table(
        doc,
        ["영역", "설명"],
        [
            ["상단 툴바", "가져오기, 내보내기, 저장, 검증 버튼 및 테이블·관계 개수 표시"],
            ["좌측 사이드바", "프로젝트 목록, 메타데이터(시스템명·작성일·작성자·DB명·스키마), 도움말"],
            ["캔버스", "테이블 노드와 관계선을 배치·편집하는 작업 영역"],
            ["하단 도구", "새 테이블, 이름 표시(한+영/한글/영문), 연결선, 관계명 표시 등"],
        ],
    )
    doc.add_paragraph()

    add_heading(doc, "2.2 프로젝트 관리", 2)
    add_bullets(
        doc,
        [
            "새 프로젝트: 사이드바 상단 「+」로 빈 프로젝트를 추가합니다.",
            "프로젝트 전환: 목록에서 프로젝트를 클릭하면 해당 작업으로 전환됩니다.",
            "저장: 「저장」 버튼 또는 자동 저장으로 브라우저 로컬 저장소에 보관됩니다.",
            "메타데이터: 시스템명, 작성일, 작성자, DB명, 스키마는 내보내기 시 반영됩니다.",
        ],
    )

    add_heading(doc, "2.3 가져오기", 2)
    add_para(doc, "상단 「가져오기」 버튼을 클릭합니다.")
    add_bullets(
        doc,
        [
            "Excel 테이블정의서: 양식에 맞는 .xlsx 파일을 선택합니다.",
            "CREATE TABLE SQL: SQL 텍스트를 붙여넣거나 파일에서 읽습니다.",
            "가져오기 모드: 「새로 만들기」(기존 작업 대체) 또는 「기존에 추가」(병합)를 선택합니다.",
        ],
    )

    add_heading(doc, "2.4 테이블·컬럼 편집", 2)
    add_numbered(
        doc,
        [
            "새 테이블: 캔버스 하단 「새 테이블」 아이콘 → 캔버스 클릭.",
            "테이블 이름 수정: 테이블 선택 후 이름 클릭 또는 더블클릭 → 수정 대화상자.",
            "컬럼 추가/수정: 컬럼 행의 편집 아이콘 → PK, FK, NULL, 타입, 한글명 등 설정.",
            "테이블 이동: 선택 도구 상태에서 테이블을 드래그합니다.",
            "다중 선택: Shift+클릭으로 여러 테이블을 선택합니다.",
            "실행 취소: Ctrl+Z로 직전에 삭제한 테이블·관계·컬럼을 복원합니다.",
        ],
    )

    add_heading(doc, "2.5 관계(연결선) 편집", 2)
    add_image(doc, "manual-er-relation-edit.png", "그림 2-2. 관계 연결 및 카디널리티 표시")
    add_numbered(
        doc,
        [
            "연결 도구: 하단 「연결선」 아이콘을 누른 뒤 첫 테이블을 클릭하고, 두 번째 테이블을 클릭하면 관계가 생성됩니다.",
            "연결 후 자동으로 선택(손바닥) 모드로 돌아갑니다.",
            "관계선 선택: 선을 클릭해 선택합니다.",
            "관계 수정: 관계선 더블클릭 → 관계 수정 대화상자(카디널리티, 식별/비식별, 관계명 등).",
            "선 끝점 조정: 세로선은 좌우, 가로선은 상하 방향으로 드래그해 경로를 조정합니다.",
            "카디널리티 표기: 선 양 끝에 1, 1..N, 0..N 등이 표시됩니다. (예: 1:N → 1 과 1..N)",
            "실선/점선: 자식 FK+PK 조합이면 식별관계(실선), 아니면 비식별(점선)으로 자동 설정됩니다.",
        ],
    )
    add_para(
        doc,
        "관계 저장 시 검증 메시지(경고·오류)는 ER 검증 모달이 아닌 「관계 수정」 대화상자 안에 표시됩니다. "
        "오류가 있으면 저장이 차단되고, 경고는 저장 후에도 확인할 수 있습니다.",
    )

    add_heading(doc, "2.6 검증", 2)
    add_para(doc, "상단 「검증」 버튼으로 전체 ER 모델을 점검합니다.")
    add_bullets(
        doc,
        [
            "타입 불일치, FK/PK 누락, 카디널리티 오류, Nullable 경고 등을 목록으로 표시합니다.",
            "컬럼 저장 시에는 해당 컬럼과 직접 관련된 항목만 표시됩니다.",
            "문제가 없으면 「문제가 없습니다.」 메시지가 나타납니다.",
        ],
    )

    add_heading(doc, "2.7 내보내기", 2)
    add_bullets(
        doc,
        [
            "Excel: 지정 양식의 테이블정의서로 저장. 전체 또는 선택한 테이블만 내보내기 가능.",
            "스크립트: PostgreSQL CREATE TABLE DDL 생성 → 미리보기 확인 후 저장.",
            "Index Key 등 Excel 양식 옵션은 내보내기 대화상자에서 설정합니다.",
        ],
    )

    add_heading(doc, "2.8 기타 단축·팁", 2)
    add_table(
        doc,
        ["기능", "조작 방법"],
        [
            ["이름 표시 전환", "한+영 → 한글 → 영문 순환 (하단 버튼)"],
            ["관계명 표시", "하단 토글로 1:1..N 등 관계명 on/off"],
            ["캔버스 잠금", "React Flow 컨트롤의 자물쇠 — 테이블 이동·연결 잠금"],
            ["빈 캔버스", "가져오기, 새 테이블, 또는 SQL로 시작"],
        ],
    )

    # 3. DB 표준 점검
    doc.add_page_break()
    add_heading(doc, "3. DB 표준 점검 도구", 1)
    add_para(
        doc,
        "설계서(테이블정의서·코드정의서)를 업로드하여 행정안전부 공통표준과 대조합니다. "
        "결과는 화면에서 확인하거나 Excel로 다운로드할 수 있습니다.",
    )
    add_image(doc, "manual-chk-db-std-ui.png", "그림 3-1. DB 표준 점검 도구 화면")

    add_heading(doc, "3.1 점검 종류", 2)
    add_table(
        doc,
        ["종류", "입력 파일", "점검 내용"],
        [
            ["표준단어", "테이블정의서 Excel", "테이블명·컬럼 한글명을 표준단어로 분해·매칭"],
            ["표준용어", "테이블정의서 Excel", "한글명이 공통표준용어와 일치하는지, 영문약어 검토"],
            ["표준도메인", "테이블정의서 Excel", "용어 매칭 후 데이터타입·길이가 표준도메인과 일치하는지"],
            ["표준코드", "코드정의서 Excel", "코드값·코드명이 표준코드 사전과 일치하는지"],
        ],
    )
    doc.add_paragraph()

    add_heading(doc, "3.2 점검 실행 절차", 2)
    add_numbered(
        doc,
        [
            "(선택) 「샘플 데이터」에서 해당 종류의 샘플 Excel을 다운로드해 형식을 확인합니다.",
            "「종류」 드롭다운에서 표준단어·표준용어·표준도메인·표준코드 중 하나를 선택합니다.",
            "표준단어/용어/도메인: 「테이블정의서 Excel」 업로드. 표준코드: 「코드정의서 Excel」 업로드.",
            "시트명: 비우면 자동 감지. 특정 시트만 점검할 경우 시트명을 입력합니다.",
            "파일 선택 후 형식 확인 메시지가 「점검 가능」이면 「점검 실행」을 클릭합니다.",
            "진행률 표시 후 결과가 하단에 나타납니다.",
        ],
    )

    add_heading(doc, "3.3 표준 CSV (선택)", 2)
    add_para(
        doc,
        "파일을 올리지 않으면 서버에 포함된 MOIS 기본 표준 사전으로 점검합니다. "
        "기관·프로젝트 전용 사전이 있을 때만 아래 CSV를 추가합니다.",
    )
    add_bullets(
        doc,
        [
            "표준단어 CSV — 표준단어·표준용어 점검, 표준용어 생성 보조",
            "표준용어 CSV — 표준용어·표준도메인 점검 및 표준용어 생성",
            "표준도메인 CSV — 표준도메인 점검 (기관별 타입·길이 정의)",
        ],
    )

    add_heading(doc, "3.4 결과 확인", 2)
    add_bullets(
        doc,
        [
            "통계 카드: 설계 컬럼 수, 일치/검토/미매칭 건수 등 요약.",
            "탭: 일치 · 검토 · 미매칭 — 각 탭별 상세 행을 표로 표시 (100건 단위 페이지).",
            "검색: 열 값으로 필터링.",
            "검토 탭: 「영문 불일치만」 체크로 영문 관련 검토 항목만 표시.",
            "결과 Excel 다운로드: 점검 결과 전체를 .xlsx로 저장.",
            "단어집/용어집 다운로드: 표준단어·표준용어 점검에서 사용된 표준 항목 및 미등록 후보.",
        ],
    )

    add_heading(doc, "3.5 표준용어 생성", 2)
    add_para(
        doc,
        "설계서 점검과 별도 기능입니다. 한글 테이블명·컬럼명 후보를 한 줄에 하나씩 입력하면 "
        "공통표준용어 조회 및 표준단어 조합 결과를 제공합니다.",
    )
    add_numbered(
        doc,
        [
            "왼쪽 「한글명」 입력란에 후보를 줄 단위로 입력합니다.",
            "「표준용어 생성」 클릭 → 오른쪽 「권장 영문명」 및 하단 표에 결과 표시.",
            "「권장 영문명 전체 복사」로 설계서 작성에 활용합니다.",
            "점검 실행 영역에서 선택한 표준단어·표준용어 CSV가 있으면 동일하게 적용됩니다.",
        ],
    )

    add_heading(doc, "3.6 DBManager 연동", 2)
    add_para(
        doc,
        "DBManager에서 설계서를 DB 표준 점검 도구로 넘기면 테이블정의서 파일이 자동으로 로드됩니다. "
        "이전 앱 이름과 파일명이 상태 메시지로 표시됩니다.",
    )

    # 4. Troubleshooting
    doc.add_page_break()
    add_heading(doc, "4. 문제 해결", 1)
    add_table(
        doc,
        ["증상", "확인 사항"],
        [
            ["API 연결 실패", "API 서버(8001) 실행 여부, NEXT_PUBLIC_API_BASE_URL 설정"],
            ["형식 확인 실패", "Excel 양식(목록형/블록형 테이블정의서, 코드정의서) 및 시트명"],
            ["점검 결과 없음", "「점검 실행」 완료 후 하단 결과 패널 확인; 종류 변경 시 캐시된 결과 재표시"],
            ["ER Modeler 저장 안 됨", "브라우저 로컬 저장소 용량·시크릿 모드 제한"],
            ["관계 저장 불가", "관계 수정 대화상자의 오류 메시지 확인 후 FK/PK·타입 수정"],
        ],
    )

    add_heading(doc, "5. 부록 — 용어", 1)
    add_table(
        doc,
        ["용어", "설명"],
        [
            ["카디널리티", "관계 양쪽의 개수 제약 (1, 1..N, 0..N 등)"],
            ["식별관계", "자식 PK에 FK가 포함된 관계 (실선)"],
            ["비식별관계", "FK가 PK가 아닌 관계 (점선)"],
            ["일치/검토/미매칭", "표준 점검 결과 등급 — 완전 일치 / 검토 필요 / 표준 미매칭"],
            ["MOIS", "행정안전부 공통표준 데이터"],
        ],
    )

    OUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUT_PATH))
    print(f"Saved: {OUT_PATH}")


if __name__ == "__main__":
    build()
